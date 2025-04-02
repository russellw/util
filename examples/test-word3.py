from docx import Document
import sys

def extract_text_from_table(table):
    text = []
    for row in table.rows:
        for cell in row.cells:
            text.append(cell.text)
            for table in cell.tables:
                text.extend(extract_text_from_table(table))
    return text

def extract_text_from_header_footer(header_footer):
    return [paragraph.text for paragraph in header_footer.paragraphs]

def convert_docx_to_txt(path):
    doc = Document(path)
    full_text = []

    # Extract text from document's paragraphs
    for para in doc.paragraphs:
        full_text.append(para.text)

    # Extract text from tables
    for table in doc.tables:
        full_text.extend(extract_text_from_table(table))

    # Extract text from headers and footers
    for section in doc.sections:
        header = section.header
        footer = section.footer

        full_text.extend(extract_text_from_header_footer(header))
        full_text.extend(extract_text_from_header_footer(footer))

    return '\n'.join(full_text)

# Path to your Word document
docx_path = sys.argv[1]
txt_content = convert_docx_to_txt(docx_path)
print(txt_content)

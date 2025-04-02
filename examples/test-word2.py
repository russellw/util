from docx import Document
import sys

def extract_text_from_table(table):
    return '\n'.join(cell.text for row in table.rows for cell in row.cells)

def extract_text_from_cell(cell):
    text = []
    for paragraph in cell.paragraphs:
        text.append(paragraph.text)
    for table in cell.tables:
        text.append(extract_text_from_table(table))
    return '\n'.join(text)

def extract_text_from_header_footer(header_footer):
    return '\n'.join(paragraph.text for paragraph in header_footer.paragraphs)

def convert_docx_to_txt(path):
    doc = Document(path)
    full_text = []

    # Extract text from document's paragraphs
    for para in doc.paragraphs:
        full_text.append(para.text)

    # Extract text from tables
    for table in doc.tables:
        full_text.append(extract_text_from_table(table))

    # Extract text from headers and footers
    for section in doc.sections:
        header = section.header
        footer = section.footer

        full_text.append(extract_text_from_header_footer(header))
        full_text.append(extract_text_from_header_footer(footer))

    # Extract text from text boxes in the document body
    for shape in doc.element.body.iter_shape():
        if shape.text:
            full_text.append(shape.text)

    return '\n'.join(full_text)

# Path to your Word document
docx_path = sys.argv[1]
txt_content = convert_docx_to_txt(docx_path)
print(txt_content)

from docx import Document
import sys

def extract_text_from_table(table):
    table_text = []
    for row in table.rows:
        for cell in row.cells:
            table_text.append(cell.text)
    return '\n'.join(table_text)

def convert_docx_to_txt(path):
    doc = Document(path)
    full_text = []

    for para in doc.paragraphs:
        full_text.append(para.text)

    for table in doc.tables:
        full_text.append(extract_text_from_table(table))
    
    return '\n'.join(full_text)

# Path to your Word document
docx_path = sys.argv[1]
txt_content = convert_docx_to_txt(docx_path)
print(txt_content)
